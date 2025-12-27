"""
Knowledge Base Service for LCP Generator.

This module handles:
1. Uploading and parsing preference documents (Word files)
2. Storing parsed preferences in Supabase
3. Retrieving preferences for use in Claude analysis

The knowledge base allows Dr. Tontz to update his clinical preferences
by simply uploading an updated Word document.
"""

import os
import json
from datetime import datetime
from typing import Optional, List, Dict
from anthropic import Anthropic

# Try to import Supabase
try:
    from supabase import create_client, Client
    SUPABASE_AVAILABLE = True
except ImportError:
    SUPABASE_AVAILABLE = False


def get_supabase_client() -> Optional['Client']:
    """Get Supabase client if configured."""
    if not SUPABASE_AVAILABLE:
        return None

    url = os.getenv('SUPABASE_URL')
    key = os.getenv('SUPABASE_KEY') or os.getenv('SUPABASE_SERVICE_KEY')

    if not url or not key:
        return None

    try:
        return create_client(url, key)
    except Exception:
        return None


def extract_text_from_docx(file_path: str) -> str:
    """Extract text content from a Word document."""
    from docx import Document

    doc = Document(file_path)
    text_content = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_content.append(paragraph.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_content.append(" | ".join(row_text))

    return "\n\n".join(text_content)


def parse_preferences_with_claude(document_text: str) -> Dict:
    """
    Use Claude to parse a preferences document and extract structured rules.

    Returns a dict with:
    - global_principles: List of general principles
    - body_part_rules: Dict of body part specific rules
    - treatment_rules: Rules about specific treatments
    - raw_summary: Human-readable summary
    """
    client = Anthropic(api_key=os.getenv('ANTHROPIC_API_KEY'))

    system_prompt = """You are an expert at parsing medical/clinical documentation.
Your task is to extract clinical decision rules and preferences from Dr. Tontz's documentation.

Extract and structure the information into these categories:

1. GLOBAL_PRINCIPLES: General rules that apply to all cases
2. SPINE_RULES: Rules specific to cervical, thoracic, or lumbar spine
3. UPPER_EXTREMITY_RULES: Rules for shoulder, elbow, wrist, hand
4. LOWER_EXTREMITY_RULES: Rules for hip, knee, ankle, foot
5. TREATMENT_RULES: Rules about specific treatments (ESI, RFA, PT, surgery, etc.)
6. AGE_RULES: Age-specific considerations
7. IMAGING_RULES: Rules about imaging surveillance

For each rule, extract:
- rule_name: Short descriptive name
- condition: When this rule applies
- action: What to do when the condition is met
- body_part: Specific body part if applicable (cervical, lumbar, shoulder, etc.)
- priority: Suggested priority (100=low, 200=medium, 300=high)

Return a JSON object with this structure:
{
    "global_principles": [
        {"rule_name": "...", "condition": "...", "action": "...", "priority": 200}
    ],
    "spine_rules": [...],
    "upper_extremity_rules": [...],
    "lower_extremity_rules": [...],
    "treatment_rules": [...],
    "age_rules": [...],
    "imaging_rules": [...],
    "raw_summary": "A 2-3 paragraph summary of the key principles in this document"
}

Be thorough - extract ALL rules and preferences mentioned in the document.
If a section doesn't have rules, return an empty array for that section."""

    user_prompt = f"""Please parse the following clinical preferences document and extract all rules and guidelines:

---
{document_text}
---

Extract all clinical decision rules into the structured JSON format."""

    try:
        response = client.messages.create(
            model="claude-sonnet-4-20250514",  # Use Sonnet for parsing (faster, cheaper)
            max_tokens=8000,
            messages=[
                {"role": "user", "content": user_prompt}
            ],
            system=system_prompt
        )

        response_text = response.content[0].text

        # Try to parse JSON from the response
        if "```json" in response_text:
            json_start = response_text.find("```json") + 7
            json_end = response_text.find("```", json_start)
            response_text = response_text[json_start:json_end].strip()
        elif "```" in response_text:
            json_start = response_text.find("```") + 3
            json_end = response_text.find("```", json_start)
            response_text = response_text[json_start:json_end].strip()

        return json.loads(response_text)

    except json.JSONDecodeError as e:
        return {
            "error": f"Failed to parse Claude response: {str(e)}",
            "raw_response": response_text if 'response_text' in locals() else "No response",
            "global_principles": [],
            "raw_summary": "Parsing failed - please try again"
        }
    except Exception as e:
        return {
            "error": f"Claude API error: {str(e)}",
            "global_principles": [],
            "raw_summary": "Parsing failed - please try again"
        }


def save_knowledge_base(
    document_name: str,
    raw_text: str,
    parsed_content: Dict,
    document_type: str = "master_preferences"
) -> Optional[Dict]:
    """
    Save parsed knowledge base content to Supabase.

    Args:
        document_name: Original filename
        raw_text: Full text of the document
        parsed_content: Structured rules extracted by Claude
        document_type: Type of document (master_preferences, update, etc.)

    Returns:
        The saved record or None if failed
    """
    client = get_supabase_client()
    if not client:
        return None

    try:
        # Save to knowledge_base table
        record = {
            'document_name': document_name,
            'document_type': document_type,
            'raw_text': raw_text[:50000],  # Limit to 50k chars
            'parsed_content': parsed_content,
            'raw_summary': parsed_content.get('raw_summary', ''),
            'is_active': True,
            'version': 1
        }

        # Check if there's an existing active record and increment version
        existing = client.table('knowledge_base') \
            .select('version') \
            .eq('document_type', document_type) \
            .eq('is_active', True) \
            .order('version', desc=True) \
            .limit(1) \
            .execute()

        if existing.data:
            # Deactivate old versions
            client.table('knowledge_base') \
                .update({'is_active': False}) \
                .eq('document_type', document_type) \
                .execute()
            record['version'] = existing.data[0]['version'] + 1

        response = client.table('knowledge_base').insert(record).execute()
        return response.data[0] if response.data else None

    except Exception as e:
        print(f"Error saving knowledge base: {e}")
        return None


def convert_parsed_to_rules(parsed_content: Dict) -> List[Dict]:
    """
    Convert parsed knowledge base content into clinical_rules format.

    This allows the parsed preferences to be automatically added
    as clinical rules that Claude will follow.
    """
    rules = []

    # Map categories to rule categories
    category_map = {
        'global_principles': ('general', None),
        'spine_rules': ('body_part', 'spine'),
        'upper_extremity_rules': ('body_part', 'upper_extremity'),
        'lower_extremity_rules': ('body_part', 'lower_extremity'),
        'treatment_rules': ('treatment_history', None),
        'age_rules': ('age', None),
        'imaging_rules': ('diagnosis', 'imaging'),
    }

    for section, (category, subcategory) in category_map.items():
        section_rules = parsed_content.get(section, [])
        if not isinstance(section_rules, list):
            continue

        for rule in section_rules:
            if not isinstance(rule, dict):
                continue

            rules.append({
                'category': category,
                'subcategory': rule.get('body_part') or subcategory,
                'rule_name': rule.get('rule_name', 'Unnamed Rule'),
                'condition_description': rule.get('condition', ''),
                'action_description': rule.get('action', ''),
                'priority': rule.get('priority', 150),
                'is_active': True,
                'created_by': 'Knowledge Base Import'
            })

    return rules


def import_rules_from_knowledge_base(parsed_content: Dict) -> int:
    """
    Import parsed rules into the clinical_rules table.

    Returns the number of rules imported.
    """
    client = get_supabase_client()
    if not client:
        return 0

    rules = convert_parsed_to_rules(parsed_content)
    if not rules:
        return 0

    try:
        # Mark existing KB-imported rules as inactive
        client.table('clinical_rules') \
            .update({'is_active': False}) \
            .eq('created_by', 'Knowledge Base Import') \
            .execute()

        # Insert new rules
        response = client.table('clinical_rules').insert(rules).execute()
        return len(response.data) if response.data else 0

    except Exception as e:
        print(f"Error importing rules: {e}")
        return 0


def get_active_knowledge_base() -> Optional[Dict]:
    """Get the currently active knowledge base content."""
    client = get_supabase_client()
    if not client:
        return None

    try:
        response = client.table('knowledge_base') \
            .select('*') \
            .eq('is_active', True) \
            .order('created_at', desc=True) \
            .limit(1) \
            .execute()

        return response.data[0] if response.data else None

    except Exception:
        return None


def get_knowledge_base_for_prompt() -> str:
    """
    Get the knowledge base content formatted for inclusion in Claude's prompt.

    This is the key function that gives Claude "memory" of the preferences.
    """
    kb = get_active_knowledge_base()
    if not kb:
        return ""

    parsed = kb.get('parsed_content', {})
    if not parsed:
        return ""

    lines = [
        "## Dr. Tontz's Clinical Preferences (from uploaded Master Document)",
        "",
        f"**Document**: {kb.get('document_name', 'Master Preferences')}",
        f"**Last Updated**: {kb.get('created_at', 'Unknown')[:10]}",
        "",
    ]

    # Add summary
    if parsed.get('raw_summary'):
        lines.append("### Summary")
        lines.append(parsed['raw_summary'])
        lines.append("")

    # Add key principles
    if parsed.get('global_principles'):
        lines.append("### Global Principles")
        for principle in parsed['global_principles'][:10]:  # Limit to 10
            if isinstance(principle, dict):
                lines.append(f"- **{principle.get('rule_name', 'Rule')}**: {principle.get('action', '')}")
        lines.append("")

    return "\n".join(lines)


def get_knowledge_base_history() -> List[Dict]:
    """Get history of all knowledge base uploads."""
    client = get_supabase_client()
    if not client:
        return []

    try:
        response = client.table('knowledge_base') \
            .select('id, document_name, document_type, version, is_active, created_at, raw_summary') \
            .order('created_at', desc=True) \
            .limit(20) \
            .execute()

        return response.data or []

    except Exception:
        return []

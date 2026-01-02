"""
AI service for medical record analysis - Phase 1.

This module handles the AI-driven clinical decision phase:
- Analyzes medical records to identify injuries/diagnoses
- Applies decision trees to match clinical scenarios
- Outputs scenario codes (C1, C4, L2, etc.)

Supports multiple AI providers:
- Google Gemini (default - fastest)
- Anthropic Claude (fallback)

The scenario codes are then mapped to item bundles by scenario_mapper.py (Phase 2).
"""

import os
import json
from app.services.scenario_bundles import get_scenario_summary
from app.services.custom_rules import get_rules_for_analysis
from app.services.knowledge_base import get_knowledge_base_for_prompt
from app.services.causation_analyzer import get_causation_protocol_for_prompt


def get_ai_provider():
    """Determine which AI provider to use based on environment variables."""
    # Check for Gemini API key first (preferred for speed)
    if os.getenv('GOOGLE_API_KEY') or os.getenv('GEMINI_API_KEY'):
        return 'gemini'
    # Fall back to Claude
    if os.getenv('ANTHROPIC_API_KEY'):
        return 'claude'
    return None


def get_scenario_list_for_prompt() -> str:
    """Generate scenario list for Claude prompt."""
    scenarios = get_scenario_summary()
    lines = []
    for code, info in sorted(scenarios.items()):
        lines.append(f"- **{code}**: {info['name']} - {info['description']}")
    return "\n".join(lines)


def _build_user_causation_section(causation_data: dict) -> str:
    """
    Build the causation protocol section using the user's uploaded causation analysis.

    This replaces AI-based causation analysis with the user's predetermined causation determinations.
    The AI should NOT re-evaluate causation - it should ACCEPT the user's determinations as final.
    """
    from app.services.causation_parser import get_causal_body_parts, get_excluded_body_parts

    causal_parts = get_causal_body_parts(causation_data)
    excluded_parts = get_excluded_body_parts(causation_data)

    causal_list = ", ".join(causal_parts) if causal_parts else "None specified"

    # Build exclusion details
    exclusion_details = []
    for entry in causation_data.get("excluded_body_parts", []):
        bp = entry.get("body_part", "Unknown")
        reason = entry.get("reason", "Not causally related")
        exclusion_details.append(f"- **{bp}**: {reason}")

    exclusion_text = "\n".join(exclusion_details) if exclusion_details else "- None specified"

    # Always include raw text so AI can see the full causation analysis
    raw_text = causation_data.get("raw_text", "")[:4000]  # First 4000 chars

    return f"""## USER-PROVIDED CAUSATION ANALYSIS - MANDATORY COMPLIANCE

###############################################################################
# CRITICAL: YOU MUST FOLLOW THESE CAUSATION DETERMINATIONS EXACTLY           #
# THE USER HAS ALREADY PERFORMED CAUSATION ANALYSIS - DO NOT OVERRIDE IT     #
###############################################################################

### Body Parts CAUSALLY RELATED (You MAY include scenarios for these):
{causal_list}

### Body Parts NOT CAUSALLY RELATED (You MUST EXCLUDE these - NO EXCEPTIONS):
{exclusion_text}

### MANDATORY RULES - VIOLATION IS NOT ACCEPTABLE:
1. **NEVER** assign scenarios to body parts listed as NOT CAUSALLY RELATED
2. **NEVER** include knee scenarios (K1-K8) if knee is listed as excluded
3. **NEVER** include shoulder scenarios (S1-S6) if shoulder is listed as excluded
4. **NEVER** include any recommendations for excluded body parts
5. If you see ANY diagnosis for an excluded body part, put it in "excluded_diagnoses" - NOT in scenarios
6. The user's causation analysis is FINAL - do not re-evaluate or override it

### Raw Causation Analysis Document:
```
{raw_text}
```

### BEFORE YOU OUTPUT:
Double-check: Did you include any scenarios for excluded body parts? If yes, REMOVE THEM.
For each scenario you output, verify the body part is NOT in the exclusion list above.
"""


def analyze_medical_records(medical_summary: str, patient_info: dict, provider_recommendations: str = "", causation_data: dict = None) -> dict:
    """
    Phase 1: Analyze medical records and identify applicable scenarios.

    Args:
        medical_summary: Text content from the medical records document
        patient_info: Dict with patient name, DOB, life expectancy, etc.
        provider_recommendations: Optional text from treating provider recommendations document
        causation_data: Optional dict from user's causation analysis document containing
                       which body parts are causally related vs excluded

    Returns:
        Dict with:
        - scenarios: List of scenario codes (e.g., ["C1", "C4"])
        - diagnoses: List of identified diagnoses with body parts
        - rationales: Dict mapping scenario code to patient-specific rationale
        - provider_items: List of items directly recommended by treating providers
        - excluded_diagnoses: List of diagnoses excluded due to causation
        - summary: Brief summary of injuries and care needs
    """
    # Determine which AI provider to use
    provider = get_ai_provider()

    if provider == 'gemini':
        return _analyze_with_gemini(medical_summary, patient_info, provider_recommendations, causation_data)
    elif provider == 'claude':
        return _analyze_with_claude(medical_summary, patient_info, provider_recommendations, causation_data)
    else:
        return {
            "error": "No AI API key configured. Set GOOGLE_API_KEY or ANTHROPIC_API_KEY.",
            "scenarios": [],
            "diagnoses": [],
            "excluded_diagnoses": [],
            "rationales": {},
            "provider_items": [],
            "summary": "Analysis failed - no API key"
        }


def _build_prompts(patient_info: dict, medical_summary: str, provider_recommendations: str = "", causation_data: dict = None):
    """Build the system and user prompts for analysis."""
    from app.services.causation_parser import get_causation_summary, get_causal_body_parts, get_excluded_body_parts

    scenario_list = get_scenario_list_for_prompt()
    custom_rules = get_rules_for_analysis()
    knowledge_base = get_knowledge_base_for_prompt()

    # If user provided causation data, use it instead of AI-based causation
    if causation_data:
        causation_protocol = _build_user_causation_section(causation_data)
    else:
        causation_protocol = get_causation_protocol_for_prompt()

    system_prompt = f"""You are a medical expert assistant helping Dr. William Tontz, MD, CLCP identify applicable clinical scenarios for Life Care Plans.

{custom_rules}

{knowledge_base}

{causation_protocol}

Your task is to analyze medical records, DETERMINE CAUSATION for each diagnosis, and identify which predefined clinical scenarios apply to this patient.

**CRITICAL CAUSATION REQUIREMENT:**
Before assigning ANY scenario code, you MUST first determine whether each diagnosis is causally related to the injury event.
Only diagnoses classified as "causal" or "aggravation" qualify for LCP recommendations.
Diagnoses classified as "exacerbation", "sprain_strain", or "not_causal" must be EXCLUDED from scenario assignment.

## Available Clinical Scenarios:
{scenario_list}

## Decision Tree for Scenario Selection:

### SPINE ANALYSIS:
For each spine region (cervical, thoracic, lumbar):

1. **Is there a STRUCTURAL diagnosis?** (herniation, stenosis, fracture, facet syndrome, radiculopathy)
   - NO structural finding → Skip this region (sprains/strains don't need LCP items)
   - YES → Continue to step 2

2. **Check for prior surgery:**
   - Cervical fusion/ACDF/disc replacement → C6
   - Thoracic fusion → T4
   - Lumbar fusion → L6
   - Lumbar discectomy → L5

3. **If no surgery, check for facet/RFA history:**
   - Facet blocks/MBB performed with benefit + NO prior RFA → One-time RFA (C4, L3)
   - Prior RFA with benefit → Annual RFA (C5, L4)

4. **If no surgery or RFA, check for radiculopathy/ESI:**
   - Active radiculopathy + NO prior ESI → One-time ESI (C2, L7)
   - Prior ESI with benefit → Annual ESI (C3, L2, T2)

5. **If structural diagnosis but none of above:**
   - Cervical disc/stenosis → C1
   - Thoracic pain → T1
   - Thoracic compression fracture → T3
   - Lumbar disc → L1

### SHOULDER ANALYSIS:
- Rotator cuff tendinopathy or partial tear, non-op → S1
- Full-thickness rotator cuff tear, non-op → S2
- Post-op rotator cuff repair → S3
- Labral tear, non-op → S4
- Post-op labral repair → S5
- Shoulder arthroplasty (replacement) → S6

### ELBOW ANALYSIS:
- Elbow tendinopathy (tennis/golfer's elbow), non-op with structural damage → E1
- Post-op elbow surgery (ulnar nerve transposition, tendon repair) → E2
- Cubital tunnel syndrome, non-surgical → E3
- Post-op cubital tunnel release → E4
- Elbow fracture/dislocation, non-op → E5
- Post-op elbow fracture ORIF → E6

### WRIST/HAND ANALYSIS:
- Wrist/hand tendinopathy or mild carpal tunnel, non-op → W1
- Post-op tendon repair of wrist/hand → W2
- Carpal tunnel syndrome, non-surgical → W3
- Post-op carpal tunnel release → W4
- Wrist fracture (distal radius, scaphoid), non-op → W5
- Post-op wrist fracture ORIF → W6

### HIP ANALYSIS:
- Hip labral tear, non-op → H1
- Post-op hip arthroscopy → H2
- Hip osteoarthritis, non-op → H3
- Hip fracture, non-op → H4
- Post-op hip fracture ORIF → H5
- Total hip arthroplasty → H6

### KNEE ANALYSIS:
- Meniscus tear, non-op → K1
- Post-op knee arthroscopy (meniscectomy/meniscal repair) → K2
- ACL tear, non-op → K3
- Post-op ACL reconstruction → K4
- Knee osteoarthritis, non-op → K5
- Tibial plateau fracture, non-op → K6
- Post-op tibial plateau fracture ORIF → K7
- Total knee arthroplasty → K8

### FOOT/ANKLE ANALYSIS:
- Ankle ligament injury with structural damage (complete tear, chronic instability), non-op → F1
- Post-op ankle ligament reconstruction → F2
- Achilles tendon injury (tendinopathy, partial/complete tear), non-op → F3
- Post-op Achilles tendon repair → F4
- Ankle fracture, non-op → F5
- Post-op ankle fracture ORIF → F6
- Ankle osteoarthritis, non-op → F7
- Ankle fusion (arthrodesis) → F8
- Total ankle arthroplasty → F9
- Foot fracture (metatarsal, calcaneus, navicular, cuboid, Lisfranc), non-op → F10
- Post-op foot fracture ORIF → F11
- Plantar fasciitis with structural fascia damage → F12

## Key Rules:
1. ONLY select scenarios for STRUCTURAL injuries (herniations, tears, fractures, stenosis)
2. Sprains/strains do NOT get scenarios - they heal in 6-12 weeks
3. Multiple scenarios can apply to the same patient (e.g., C1 + C4 for disc + facet)
4. ONE-TIME vs RECURRING is determined by treatment history (see decision trees above)
5. Rationales must be based purely on medical record findings - NO mention of scenario codes or system

## Treating Provider Recommendations:
If treating provider recommendations are included, you MUST:
1. Extract SPECIFIC recommendations with exact frequencies stated by the provider
2. Identify the provider's name and credentials
3. Quote the provider's recommendation verbatim when possible
4. Create provider_items for any recommendation that doesn't fit into a standard scenario
5. Suggest the most appropriate CPT code for each provider-recommended item

Provider recommendations should be cited in rationales like:
- "Dr. Smith recommended MRI of the cervical spine every 2 years for the duration of life expectancy."
- "Per Dr. Johnson's 7/15/25 recommendations, the patient will require annual EMG/NCS studies."

## Common CPT Codes Reference (for provider items):
OFFICE/OUTPATIENT VISITS:
- 99213: Office visit, established patient, low complexity
- 99214: Office visit, established patient, moderate complexity
- 99215: Office visit, established patient, high complexity
- 99203: Office visit, new patient, low complexity
- 99204: Office visit, new patient, moderate complexity
- 99205: Office visit, new patient, high complexity

IMAGING:
- 72141: MRI cervical spine without contrast
- 72146: MRI thoracic spine without contrast
- 72148: MRI lumbar spine without contrast
- 73221: MRI shoulder without contrast
- 73721: MRI hip without contrast
- 73721: MRI knee without contrast
- 72170: X-ray pelvis
- 73030: X-ray shoulder
- 73560: X-ray knee

DIAGNOSTIC TESTING:
- 95885: EMG/NCS limited study
- 95886: EMG/NCS complete study
- 95910: Nerve conduction studies, 7-8 studies
- 95911: Nerve conduction studies, 9-10 studies

SPECIALTY CONSULTATIONS:
- 99243: Office consultation, low complexity
- 99244: Office consultation, moderate complexity
- 99245: Office consultation, high complexity

THERAPY:
- 97110: Therapeutic exercises
- 97140: Manual therapy
- 97530: Therapeutic activities
- 97542: Wheelchair management training

## Output Format:
Return a JSON object with this structure:
{{
    "scenarios": ["C1", "C4"],
    "diagnoses": [
        {{
            "body_part": "Cervical Spine",
            "diagnosis": "C5-6 disc herniation",
            "structural": true,
            "date_documented": "7/10/2025",
            "causation": "causal",
            "causation_rationale": "New cervical radiculopathy documented in ED on 1/15/25, 2 days post-MVA. No prior cervical complaints in 5 years of records reviewed. Mechanism consistent with disc injury."
        }},
        {{
            "body_part": "Cervical Spine",
            "diagnosis": "Facet syndrome",
            "structural": true,
            "date_documented": "7/10/2025",
            "causation": "causal",
            "causation_rationale": "Facet-mediated pain documented 3 weeks post-MVA after extension/rotation mechanism. Diagnostic MBB on 8/15/25 confirmed facet origin. No prior facet treatment."
        }},
        {{
            "body_part": "Lumbar Spine",
            "diagnosis": "L4-5 stenosis",
            "structural": true,
            "date_documented": "7/10/2025",
            "causation": "not_causal",
            "causation_rationale": "Pre-existing stenosis documented on 2018 MRI. No change in symptoms or treatment pattern post-accident. Excluded from LCP."
        }}
    ],
    "excluded_diagnoses": [
        {{
            "body_part": "Lumbar Spine",
            "diagnosis": "L4-5 stenosis",
            "causation": "not_causal",
            "reason": "Pre-existing condition without aggravation - no change in care pattern"
        }},
        {{
            "body_part": "Right Shoulder",
            "diagnosis": "Rotator cuff strain",
            "causation": "sprain_strain",
            "reason": "Muscular strain expected to resolve in 6-12 weeks - no ongoing LCP needs"
        }}
    ],
    "rationales": {{
        "C1": "Annual surveillance of C5-6 disc herniation with moderate foraminal stenosis per 7/10/25 MRI.",
        "C4": "One-time cervical RFA for facet-mediated pain with documented 80% relief from 8/15/25 medial branch blocks."
    }},
    "provider_items": [
        {{
            "item": "MRI Cervical Spine",
            "frequency": "Every 2 years",
            "provider_name": "Dr. John Smith",
            "provider_quote": "The patient will require MRI of the cervical spine every 2 years to monitor disc progression.",
            "body_part": "Cervical Spine",
            "suggested_cpt": "72141",
            "suggested_category": "Diagnostic Testing/Assessment",
            "rationale": "Dr. John Smith recommended MRI of the cervical spine every 2 years for the duration of life expectancy to monitor disc progression."
        }},
        {{
            "item": "Otolaryngology Follow-up",
            "frequency": "Yearly",
            "provider_name": "Dr. Jane Doe",
            "provider_quote": "Annual ENT follow-up recommended for ongoing management.",
            "body_part": "Head/Neck",
            "suggested_cpt": "99214",
            "suggested_category": "Physician/Nurse Evaluations",
            "rationale": "Dr. Jane Doe recommended annual otolaryngology follow-up for ongoing management."
        }}
    ],
    "summary": "52-year-old with cervical disc herniation and facet syndrome causally related to 1/13/25 MVA. Pre-existing lumbar stenosis excluded (no aggravation). MBB provided significant relief."
}}

CRITICAL:
- **CAUSATION FIRST**: You MUST determine causation for EVERY diagnosis BEFORE assigning scenarios
- Only assign scenarios to diagnoses with causation = "causal" or "aggravation"
- EXCLUDE diagnoses with causation = "exacerbation", "sprain_strain", or "not_causal" from scenarios
- Include ALL excluded diagnoses in the "excluded_diagnoses" array with clear reasons
- Scenarios must be from the available list above
- Rationales must reference SPECIFIC findings and DATES from the medical records
- Rationales should read naturally - DO NOT mention scenario codes (C1, C4, etc.) or "clinical scenario"
- Be conservative - typical cases have 2-5 scenarios, NOT 10+
- Provider recommendations take PRECEDENCE - if a provider specifies a frequency, use that exact frequency
- Always include the provider's name when citing their recommendation
- Sprain/strain CANNOT coexist with deeper diagnoses (disc herniation, stenosis, fracture, labral/meniscal tear) in the same body part
"""

    # Build the user prompt with optional provider recommendations section
    provider_section = ""
    if provider_recommendations:
        provider_section = f"""

## Treating Provider Recommendations
{provider_recommendations}

"""

    user_prompt = f"""## Patient Information
- Name: {patient_info.get('patient_name', 'Unknown')}
- Date of Birth: {patient_info.get('date_of_birth', 'Unknown')}
- Date of Injury: {patient_info.get('date_of_injury', 'Unknown')}
- Life Expectancy: {patient_info.get('life_expectancy', 'Unknown')} years
- Age at Report: {patient_info.get('age_initiated', 'Unknown')}

## Medical Records Summary
{medical_summary if medical_summary else "No medical summary provided."}
{provider_section}
Please analyze the medical records above and:
1. Identify all structural injuries by body region
2. Apply the decision trees to determine which scenarios apply
3. Provide patient-specific rationales for each scenario
4. {"Extract ALL treating provider recommendations with exact frequencies and provider names" if provider_recommendations else ""}
5. Return your analysis as a JSON object

Remember: Only include scenarios for STRUCTURAL injuries. Sprains/strains get NO scenarios.
{"IMPORTANT: Provider recommendations are included. Extract and cite each provider recommendation with their name and exact frequency." if provider_recommendations else ""}"""

    return system_prompt, user_prompt


def _parse_json_response(response_text: str, causation_data: dict = None) -> dict:
    """Parse JSON from AI response, handling various formats."""
    json_text = response_text

    # First try: extract from ```json ... ``` blocks
    if "```json" in response_text:
        json_start = response_text.find("```json") + 7
        json_end = response_text.find("```", json_start)
        if json_end > json_start:
            json_text = response_text[json_start:json_end].strip()
    # Second try: extract from ``` ... ``` blocks
    elif "```" in response_text:
        json_start = response_text.find("```") + 3
        json_end = response_text.find("```", json_start)
        if json_end > json_start:
            json_text = response_text[json_start:json_end].strip()
    # Third try: find JSON object directly (starts with { ends with })
    elif "{" in response_text:
        first_brace = response_text.find("{")
        last_brace = response_text.rfind("}")
        if first_brace >= 0 and last_brace > first_brace:
            json_text = response_text[first_brace:last_brace + 1]

    result = json.loads(json_text)

    scenarios = result.get("scenarios", [])
    rationales = result.get("rationales", {})
    excluded_diagnoses = result.get("excluded_diagnoses", [])

    # SAFETY NET: If causation data provided, filter out any scenarios for excluded body parts
    if causation_data:
        scenarios, rationales, additional_excluded = _filter_excluded_scenarios(
            scenarios, rationales, causation_data
        )
        # Add any AI-included scenarios that should have been excluded
        excluded_diagnoses.extend(additional_excluded)

    return {
        "scenarios": scenarios,
        "diagnoses": result.get("diagnoses", []),
        "excluded_diagnoses": excluded_diagnoses,
        "rationales": rationales,
        "provider_items": result.get("provider_items", []),
        "summary": result.get("summary", ""),
        "error": None
    }


def _filter_excluded_scenarios(scenarios: list, rationales: dict, causation_data: dict) -> tuple:
    """
    Safety net: Remove any scenarios for body parts that should be excluded.
    Returns (filtered_scenarios, filtered_rationales, additional_excluded_diagnoses)
    """
    from app.services.causation_parser import get_excluded_body_parts

    excluded_body_parts = [bp.lower() for bp in get_excluded_body_parts(causation_data)]

    # Map scenario prefixes to body parts
    scenario_body_map = {
        'C': ['cervical', 'cervical spine'],
        'T': ['thoracic', 'thoracic spine'],
        'L': ['lumbar', 'lumbar spine'],
        'S': ['shoulder', 'right shoulder', 'left shoulder'],
        'E': ['elbow', 'right elbow', 'left elbow'],
        'W': ['wrist', 'right wrist', 'left wrist', 'hand', 'right hand', 'left hand'],
        'H': ['hip', 'right hip', 'left hip'],
        'K': ['knee', 'right knee', 'left knee'],
        'F': ['ankle', 'right ankle', 'left ankle', 'foot', 'right foot', 'left foot'],
    }

    filtered_scenarios = []
    filtered_rationales = {}
    additional_excluded = []

    for scenario in scenarios:
        prefix = scenario[0] if scenario else ''
        body_parts_for_scenario = scenario_body_map.get(prefix, [])

        # Check if any of this scenario's body parts are excluded
        is_excluded = any(
            bp in excluded_body_parts or any(ebp in bp for ebp in excluded_body_parts)
            for bp in body_parts_for_scenario
        )

        if is_excluded:
            # This scenario should not have been included - add to excluded list
            additional_excluded.append({
                "body_part": body_parts_for_scenario[0].title() if body_parts_for_scenario else "Unknown",
                "diagnosis": f"Scenario {scenario}",
                "causation": "not_causal",
                "reason": "Body part excluded per user's causation analysis - removed by safety filter"
            })
        else:
            filtered_scenarios.append(scenario)
            if scenario in rationales:
                filtered_rationales[scenario] = rationales[scenario]

    return filtered_scenarios, filtered_rationales, additional_excluded


def _analyze_with_gemini(medical_summary: str, patient_info: dict, provider_recommendations: str = "", causation_data: dict = None) -> dict:
    """Analyze medical records using Google Gemini."""
    try:
        import google.generativeai as genai

        # Configure Gemini
        api_key = os.getenv('GOOGLE_API_KEY') or os.getenv('GEMINI_API_KEY')
        genai.configure(api_key=api_key)

        # Build prompts (with causation data if provided)
        system_prompt, user_prompt = _build_prompts(patient_info, medical_summary, provider_recommendations, causation_data)

        # Combine system and user prompts for Gemini
        full_prompt = f"{system_prompt}\n\n---\n\n{user_prompt}"

        # Create model and generate
        model = genai.GenerativeModel('gemini-2.0-flash-exp')
        response = model.generate_content(full_prompt)

        response_text = response.text
        return _parse_json_response(response_text, causation_data)

    except json.JSONDecodeError as e:
        return {
            "error": f"Failed to parse Gemini response: {str(e)}",
            "raw_response": response_text if 'response_text' in locals() else "No response",
            "scenarios": [],
            "diagnoses": [],
            "excluded_diagnoses": [],
            "rationales": {},
            "provider_items": [],
            "summary": "Analysis failed - please try again"
        }
    except Exception as e:
        return {
            "error": f"Gemini API error: {str(e)}",
            "scenarios": [],
            "diagnoses": [],
            "excluded_diagnoses": [],
            "rationales": {},
            "provider_items": [],
            "summary": "Analysis failed - please try again"
        }


def _analyze_with_claude(medical_summary: str, patient_info: dict, provider_recommendations: str = "", causation_data: dict = None) -> dict:
    """Analyze medical records using Anthropic Claude."""
    try:
        from anthropic import Anthropic

        client = Anthropic(api_key=os.getenv('ANTHROPIC_API_KEY'))

        # Build prompts (with causation data if provided)
        system_prompt, user_prompt = _build_prompts(patient_info, medical_summary, provider_recommendations, causation_data)

        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4096,
            messages=[
                {"role": "user", "content": user_prompt}
            ],
            system=system_prompt
        )

        response_text = response.content[0].text
        return _parse_json_response(response_text, causation_data)

    except json.JSONDecodeError as e:
        return {
            "error": f"Failed to parse Claude response: {str(e)}",
            "raw_response": response_text if 'response_text' in locals() else "No response",
            "scenarios": [],
            "diagnoses": [],
            "excluded_diagnoses": [],
            "rationales": {},
            "provider_items": [],
            "summary": "Analysis failed - please try again"
        }
    except Exception as e:
        return {
            "error": f"Claude API error: {str(e)}",
            "scenarios": [],
            "diagnoses": [],
            "excluded_diagnoses": [],
            "rationales": {},
            "provider_items": [],
            "summary": "Analysis failed - please try again"
        }


def extract_text_from_docx(file_path: str) -> str:
    """Extract text content from a Word document."""
    from docx import Document

    doc = Document(file_path)
    text_content = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_content.append(paragraph.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_content.append(" | ".join(row_text))

    return "\n\n".join(text_content)
